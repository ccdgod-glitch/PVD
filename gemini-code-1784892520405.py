import streamlit as st
import numpy as np
import pandas as pd
import plotly.graph_objects as go
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------------------------------------------
# 1. PAGE CONFIGURATION & CUSTOM CSS
# ---------------------------------------------------------
st.set_page_config(
    page_title="PVD Settlement Analytics",
    page_icon="🏗️",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
    /* 1. ปรับพื้นหลังหลัก หัวข้อ และตัวหนังสือทุกประเภทบนหน้าเว็บ */
    html, body, [class*="st-"], .stMarkdown, p, span, label, div, h1, h2, h3, h4, h5, h6 {
        color: #111827 !important;
        font-family: 'Sarabun', 'Segoe UI', sans-serif !important;
    }
    .stApp, [data-testid="stAppViewContainer"], [data-testid="stHeader"] {
        background-color: #FFFFFF !important;
    }

    /* 2. แถบเมนูด้านข้าง (Sidebar) */
    [data-testid="stSidebar"], [data-testid="stSidebar"] > div {
        background-color: #F8FAFC !important;
        border-right: 2px solid #E2E8F0 !important;
    }

    /* 3. ปุ่มกดทั้งหมด (เช่น ปุ่มดาวน์โหลด CSV / Word) ให้เป็นสีฟ้า PVD เด่นชัด */
    button[kind="primary"], button[kind="secondary"], .stButton > button, .stDownloadButton > button {
        background-color: #008BAA !important;
        color: #FFFFFF !important;
        border: none !important;
        border-radius: 8px !important;
        font-weight: 600 !important;
        padding: 8px 16px !important;
        transition: all 0.2s ease !important;
    }
    button[kind="primary"]:hover, button[kind="secondary"]:hover, .stButton > button:hover, .stDownloadButton > button:hover {
        background-color: #006B83 !important;
        box-shadow: 0 4px 10px rgba(0, 139, 170, 0.25) !important;
    }

    /* 4. ตัวเลือก Radio และ Checkbox */
    [data-testid="stRadioButton"] label, [data-testid="stCheckbox"] label {
        color: #111827 !important;
        font-weight: 500 !important;
    }

    /* 5. กล่องแจ้งเตือน (st.info, st.success, st.warning) ให้เป็นสีสว่าง อ่านง่าย */
    div[data-testid="stAlertContainer"] {
        background-color: #F8FAFC !important;
        border-radius: 8px !important;
        border: 1px solid #CBD5E1 !important;
    }
    div[data-testid="stAlertContainer"] * {
        color: #111827 !important;
    }

    /* 6. สมการคณิตศาสตร์ (LaTeX / KaTeX) ตัวหนังสือสีดำเข้ม */
    .katex, .katex-display, .katex * {
        color: #111827 !important;
    }

    /* 7. ช่องกรอกข้อมูล (Selectbox, Input, NumberInput) */
    div[data-baseweb="select"] > div,
    div[data-baseweb="input"] > div,
    input, select, textarea {
        background-color: #FFFFFF !important;
        color: #111827 !important;
        border: 1.5px solid #94A3B8 !important;
        border-radius: 6px !important;
    }
    button[data-testid="stNumberInputStepDown"], 
    button[data-testid="stNumberInputStepUp"] {
        background-color: #F1F5F9 !important;
        color: #111827 !important;
        border: 1px solid #CBD5E1 !important;
    }

    /* 8. ตารางข้อมูล (ทั้ง st.table และ st.dataframe) */
    div[data-testid="stTable"] table, [data-testid="stDataFrame"] {
        background-color: #FFFFFF !important;
        border: 1px solid #CBD5E1 !important;
        border-radius: 8px !important;
    }
    div[data-testid="stTable"] th {
        background-color: #F1F5F9 !important;
        color: #0F172A !important;
        font-weight: 700 !important;
        border-bottom: 2px solid #008BAA !important;
    }
    div[data-testid="stTable"] td {
        color: #111827 !important;
        border-bottom: 1px solid #E2E8F0 !important;
    }

    /* 9. การ์ดตัวเลขสรุป (Metric Cards) แถบแดงสไลด์ */
    div[data-testid="stMetric"] {
        background-color: #FFFFFF !important;
        border-left: 5px solid #FF5964 !important;
        border-top: 1px solid #E2E8F0 !important;
        border-right: 1px solid #E2E8F0 !important;
        border-bottom: 1px solid #E2E8F0 !important;
        border-radius: 8px !important;
        box-shadow: 0 2px 5px rgba(0,0,0,0.04) !important;
    }
    div[data-testid="stMetricLabel"] > label {
        color: #475569 !important;
        font-weight: 600 !important;
    }
    div[data-testid="stMetricValue"] { 
        color: #111827 !important;
        font-weight: 800 !important;
    }

    /* 10. กล่องพับ (Expander) และแท็บ (Tabs) */
    div[data-testid="stExpander"] {
        background-color: #FFFFFF !important;
        border: 1px solid #CBD5E1 !important;
        border-radius: 10px !important;
        margin-bottom: 12px !important;
    }
    div[data-testid="stExpander"] details summary {
        background-color: #F1F5F9 !important;
        color: #008BAA !important;
        font-weight: 700 !important;
    }
    .stTabs [data-baseweb="tab-list"] { 
        background-color: #F8FAFC !important;
        border: 1px solid #E2E8F0 !important;
    }
    .stTabs [aria-selected="true"] {
        background-color: #FFFFFF !important;
        color: #008BAA !important;
        border-bottom: 3px solid #008BAA !important;
        font-weight: 700 !important;
    }
</style>
""", unsafe_allow_html=True)

# ---------------------------------------------------------
# 2. HELPER FUNCTION: สรุปผล + สร้างตารางใน WORD REPORT
# ---------------------------------------------------------
def generate_word_report(pattern, S, H_soil, H_pvd, Cv, Cr, target_day, Uav_pct, Ur_pct, Uv_pct, S_final, S_target, days_90, d_w_mm, d_e_m, n, Fn, df):
    doc = Document()
    
    # 1. หัวข้อรายงาน
    title = doc.add_heading('รายงานสรุปผลการวิเคราะห์การเร่งการทรุดตัวด้วย PVD', level=1)
    if title.runs:
        title.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    doc.add_paragraph(f"วันที่ออกรายงาน: {pd.Timestamp.now().strftime('%d/%m/%Y')}")
    doc.add_paragraph("-" * 55)
    
    # 2. ข้อมูลพารามิเตอร์การออกแบบ
    doc.add_heading('1. ข้อมูลและพารามิเตอร์การออกแบบ (Input Parameters)', level=2)
    p1 = doc.add_paragraph()
    p1.add_run(f"ทำการติดตั้ง PVD ในชั้นดินอ่อนความหนา {H_soil:.1f} m โดยจัดวางในรูปแบบ {pattern} ที่ระยะห่าง S = {S:.2f} m รายละเอียดพารามิเตอร์สรุปดังตารางด้านล่าง:")

    # --- WORD TABLE 1: Input Parameters Table ---
    t1 = doc.add_table(rows=1, cols=3)
    t1.style = 'Table Grid'
    hdr_cells1 = t1.rows[0].cells
    hdr_cells1[0].text = 'พารามิเตอร์ (Parameter)'
    hdr_cells1[1].text = 'ค่าที่ใช้ (Value)'
    hdr_cells1[2].text = 'หน่วย (Unit)'
    
    for cell in hdr_cells1:
        cell.paragraphs[0].runs[0].font.bold = True
        
    params_data = [
        ("รูปแบบการจัดวาง (Pattern)", f"{pattern}", "-"),
        ("ระยะห่างการติดตั้ง (S)", f"{S:.2f}", "m"),
        ("ความหนาชั้นดินอ่อน (H_soil)", f"{H_soil:.1f}", "m"),
        ("ความลึกแผ่น PVD (H_pvd)", f"{H_pvd:.1f}", "m"),
        ("ค่าสัมประสิทธิ์ Cv", f"{Cv:.1f}", "cm²/day"),
        ("ค่าสัมประสิทธิ์ Cr", f"{Cr:.1f}", "cm²/day"),
        ("ขนาดเส้นผ่านศูนย์กลาง dw / de", f"{d_w_mm:.1f} mm / {d_e_m:.2f} m", "-"),
        ("อัตราส่วน n (de/dw)", f"{n:.2f}", "-"),
        ("Drain Spacing Factor F(n)", f"{Fn:.3f}", "-"),
    ]
    
    for p_name, p_val, p_unit in params_data:
        row_cells = t1.add_row().cells
        row_cells[0].text = p_name
        row_cells[1].text = p_val
        row_cells[2].text = p_unit
        
    doc.add_paragraph() # เว้นบรรทัด

    # 3. ผลการคำนวณ ณ วันเป้าหมาย
    doc.add_heading('2. ผลการคำนวณ ณ วันเป้าหมาย (Target Day Results)', level=2)
    status_text = "บรรลุตามข้อกำหนดการออกแบบ (>= 90%)" if Uav_pct >= 90 else "ยังไม่บรรลุตามเป้าหมาย (< 90%)"
    
    p2 = doc.add_paragraph()
    p2.add_run(f"ณ ระยะเวลาเป้าหมายที่ ")
    p2.add_run(f"{target_day} วัน ").bold = True
    p2.add_run("ผลการคำนวณอัตราการอัดตัวคายน้ำและการทรุดตัว สรุปได้ดังตารางด้านล่าง:")

    # --- WORD TABLE 2: Target Day Summary Table ---
    t2 = doc.add_table(rows=1, cols=3)
    t2.style = 'Table Grid'
    hdr_cells2 = t2.rows[0].cells
    hdr_cells2[0].text = 'รายการคำนวณ'
    hdr_cells2[1].text = 'ผลการคำนวณ'
    hdr_cells2[2].text = 'หมายเหตุ / เกณฑ์'

    for cell in hdr_cells2:
        cell.paragraphs[0].runs[0].font.bold = True

    results_data = [
        ("Degree of Consolidation แนวดิ่ง (Uv)", f"{Uv_pct:.2f}%", "Terzaghi Theory"),
        ("Degree of Consolidation แนวรัศมี (Ur)", f"{Ur_pct:.2f}%", "Barron / Hansbo"),
        ("Degree of Consolidation เฉลี่ยรวม (Uav)", f"{Uav_pct:.2f}%", status_text),
        ("การทรุดตัวที่เกิดขึ้นแล้ว (St)", f"{S_target:.3f} m", f"จาก S_final = {S_final:.3f} m"),
        ("เวลาบรรลุ U = 90%", f"{days_90} วัน", "เป้าหมายมาตรฐาน")
    ]

    for r_name, r_val, r_note in results_data:
        row_cells = t2.add_row().cells
        row_cells[0].text = r_name
        row_cells[1].text = r_val
        row_cells[2].text = r_note

    doc.add_paragraph() # เว้นบรรทัด

    # 4. ตารางพัฒนาการตามช่วงเวลา (Progress Table)
    doc.add_heading('3. ตารางพัฒนาการอัดตัวคายน้ำและการทรุดตัวตามช่วงเวลา', level=2)
    
    # --- WORD TABLE 3: Consolidation Progress Over Time Table ---
    t3 = doc.add_table(rows=1, cols=6)
    t3.style = 'Table Grid'
    hdr_cells3 = t3.rows[0].cells
    hdr_cells3[0].text = 'เวลา t (วัน)'
    hdr_cells3[1].text = 'U_r (%)'
    hdr_cells3[2].text = 'U_v (%)'
    hdr_cells3[3].text = 'U_av (%)'
    hdr_cells3[4].text = 'การทรุดตัว S_t (m)'
    hdr_cells3[5].text = 'สถานะ (> 90%)'

    for cell in hdr_cells3:
        cell.paragraphs[0].runs[0].font.bold = True

    test_days = sorted(list(set([30, 60, target_day, 180, 270, 365])))
    for d_val in test_days:
        if d_val <= 365:
            r = df[df["Day"] == d_val].iloc[0]
            ur_p, uv_p, uav_p, st_p = r['U_r'], r['U_v'], r['U_av'], r['Settlement']
            st_status = "ผ่านเกณฑ์" if uav_p >= 90 else "ยังไม่ถึงเกณฑ์"
        else:
            ur_p = (1 - np.exp((-8 * (Cr * d_val) / (d_e_m * 100)**2) / Fn)) * 100
            tv_tmp = (Cv * d_val) / ((H_soil*100/2)**2)
            uv_tmp = np.sqrt(4 * tv_tmp) / np.pi if tv_tmp <= 0.286 else 1 - (10**(-0.085 - 0.933 * tv_tmp))
            uv_p = min(uv_tmp * 100, 100)
            uav_p = 100 * (1 - (1 - ur_p/100) * (1 - uv_p/100))
            st_p = (uav_p / 100.0) * S_final
            st_status = "ผ่านเกณฑ์" if uav_p >= 90 else "ยังไม่ถึงเกณฑ์"

        row_cells = t3.add_row().cells
        row_cells[0].text = f"{d_val} วัน"
        row_cells[1].text = f"{ur_p:.2f}%"
        row_cells[2].text = f"{uv_p:.2f}%"
        row_cells[3].text = f"{uav_p:.2f}%"
        row_cells[4].text = f"{st_p:.3f} m"
        row_cells[5].text = st_status

    doc.add_paragraph() # เว้นบรรทัด

    # 5. สรุปผลเชิงวิศวกรรม
    doc.add_heading('4. สรุปผลและข้อเสนอแนะเชิงวิศวกรรม', level=2)
    p3 = doc.add_paragraph()
    
    if Uav_pct >= 90.0:
        p3.add_run(
            f"สรุปการประเมิน: การติดตั้ง PVD ระยะห่าง {S:.2f} เมตร มีความเหมาะสมและเพียงพอทางการวิศวกรรม "
            f"โดยชั้นดินสามารถคายน้ำและเกิดการทรุดตัวได้ถึง {Uav_pct:.2f}% ภายในระยะเวลา {target_day} วัน "
            f"(ใช้เวลาเพียง {days_90} วันในการบรรลุ U=90%) ซึ่งเป็นไปตามมาตรฐานการควบคุมการทรุดตัวก่อนการเปิดใช้งานพื้นที่"
        )
    else:
        p3.add_run(
            f"ข้อสังเกตทางการวิศวกรรม: การติดตั้ง PVD ที่ระยะห่าง {S:.2f} เมตร ณ เวลา {target_day} วัน "
            f"ทำให้เกิดการอัดตัวคายน้ำได้เพียง {Uav_pct:.2f}% ซึ่งยังไม่ถึงเกณฑ์ 90% "
            f"จึงมีข้อเสนอแนะให้ปรับลดระยะห่างการติดตั้ง PVD (S) ให้แคบลง เช่น ปรับลดเหลือ {max(0.6, S-0.2):.2f} เมตร "
            f"หรือขยายระยะเวลาในการรอคอยการทรุดตัว (Preloading Period) เพิ่มเติม"
        )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ---------------------------------------------------------
# 3. HEADER
# ---------------------------------------------------------
st.markdown('<div class="main-header">🏗️ PVD Consolidation Analytics</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-header">ระบบจำลองการทรุดตัวและคำนวณเวลาการอัดตัวคายน้ำชั้นดินอ่อน (ปรับหน่วยตามตำราเรียน / สไลด์เลคเชอร์)</div>', unsafe_allow_html=True)

# ---------------------------------------------------------
# 4. SIDEBAR - INPUT PARAMETERS
# ---------------------------------------------------------
st.sidebar.title("⚙️ ตั้งค่าพารามิเตอร์การออกแบบ")

with st.sidebar.expander("📐 1. รูปแบบ PVD & ระยะติดตั้ง", expanded=True):
    pattern = st.selectbox("รูปแบบการจัดวาง", ["สามเหลี่ยม (Triangular)", "สี่เหลี่ยม (Square)"])
    S = st.number_input("ระยะห่างการติดตั้ง S (m)", value=1.00, step=0.05, format="%.2f")
    H_pvd = st.number_input("ความลึกแผ่น PVD (m)", value=10.0, step=0.5)
    a_mm = st.number_input("ความกว้าง PVD: a (mm)", value=100.0, step=5.0)
    b_mm = st.number_input("ความหนา PVD: b (mm)", value=5.0, step=0.5)

with st.sidebar.expander("🧪 2. คุณสมบัติชั้นดิน & ทางระบายน้ำ", expanded=True):
    H_soil = st.number_input("ความหนาชั้นดินอ่อน: H_soil (m)", value=30.0, step=1.0)
    drainage_type = st.radio("ลักษณะการระบายน้ำแนวดิ่ง (Drainage Path)", ["ระบายน้ำ 2 ทาง (บน-ล่าง: Hd = H/2)", "ระบายน้ำทางเดียว (Hd = H)"])
    
    st.markdown("---")
    unit_choice = st.selectbox("เลือกหน่วยของ Cv และ Cr", ["cm²/day (ตรงตามสไลด์)", "m²/year"])
    
    if "cm²/day" in unit_choice:
        Cv_input = st.number_input("ค่า Cv (cm²/day)", value=20.0, step=1.0)
        Cr_input = st.number_input("ค่า Cr (cm²/day)", value=140.0, step=5.0)
        Cv_m2_yr = (Cv_input / 10000.0) * 365.25
        Cr_m2_yr = (Cr_input / 10000.0) * 365.25
        Cv_cm2_day = Cv_input
        Cr_cm2_day = Cr_input
    else:
        Cv_m2_yr = st.number_input("ค่า Cv (m²/year)", value=2.0, step=0.1)
        Cr_m2_yr = st.number_input("ค่า Cr (m²/year)", value=6.0, step=0.1)
        Cv_cm2_day = (Cv_m2_yr / 365.25) * 10000.0
        Cr_cm2_day = (Cr_m2_yr / 365.25) * 10000.0

    st.markdown("---")
    Cc = st.number_input("Compression Index (Cc)", value=0.80, step=0.05)
    e0 = st.number_input("Initial Void Ratio (e0)", value=2.00, step=0.1)
    sigma_0 = st.number_input("Effective Stress เดิม: σ0' (kPa)", value=50.0, step=5.0)
    delta_sigma = st.number_input("น้ำหนักถมเพิ่ม: Δσ (kPa)", value=80.0, step=5.0)

with st.sidebar.expander("🎯 3. กำหนดวันเป้าหมาย & Smear Effect", expanded=True):
    target_day = st.number_input("วันเป้าหมายหลักที่ต้องการตรวจสอบ (วัน)", value=90, step=10, min_value=1)
    include_smear = st.checkbox("คิดผลกระทบ Smear Effect", value=False)
    if include_smear:
        d_s_ratio = st.number_input("อัตราส่วน ds / dw", value=2.5, step=0.1)
        kh_ks_ratio = st.number_input("อัตราส่วน kh / ks", value=3.0, step=0.5)
    else:
        d_s_ratio, kh_ks_ratio = 1.0, 1.0

# ---------------------------------------------------------
# 5. CALCULATION ENGINE
# ---------------------------------------------------------
a_m, b_m = a_mm / 1000.0, b_mm / 1000.0
d_w_m = (a_m + b_m) / 2.0  # สูตร (a+b)/2 ตรงตามสไลด์
d_w_cm = d_w_m * 100.0
d_w_mm = d_w_m * 1000.0

d_e_m = 1.05 * S if "สามเหลี่ยม" in pattern else 1.13 * S
d_e_cm = d_e_m * 100.0

n = d_e_m / d_w_m

if include_smear:
    s = d_s_ratio
    Fn = np.log(n / s) + (kh_ks_ratio * np.log(s)) - 0.75
else:
    Fn = (n**2 / (n**2 - 1)) * np.log(n) - (3 * n**2 - 1) / (4 * n**2)

H_d_m = H_soil / 2.0 if "2 ทาง" in drainage_type else H_soil
H_d_cm = H_d_m * 100.0

S_final = H_soil * (Cc / (1 + e0)) * np.log10((sigma_0 + delta_sigma) / sigma_0)

days = np.arange(1, 366)
times_years = days / 365.25
U_v_list, U_r_list, U_av_list, S_t_list = [], [], [], []

for t_yr, d_day in zip(times_years, days):
    Tv = (Cv_cm2_day * d_day) / (H_d_cm**2)
    U_v = np.sqrt(4 * Tv) / np.pi if Tv <= 0.286 else 1 - (10**(-0.085 - 0.933 * Tv))
    U_v = min(U_v, 1.0)
    
    Tr = (Cr_cm2_day * d_day) / (d_e_cm**2)
    U_r = 1 - np.exp((-8 * Tr) / Fn)
    U_r = min(U_r, 1.0)
    
    U_av = 1 - (1 - U_r) * (1 - U_v)
    S_t = U_av * S_final
    
    U_v_list.append(U_v * 100)
    U_r_list.append(U_r * 100)
    U_av_list.append(U_av * 100)
    S_t_list.append(S_t)

df = pd.DataFrame({"Day": days, "U_v": U_v_list, "U_r": U_r_list, "U_av": U_av_list, "Settlement": S_t_list})

u90_idx = df[df["U_av"] >= 90].first_valid_index()
days_90 = df.loc[u90_idx, "Day"] if u90_idx is not None else "> 365"

# ---------------------------------------------------------
# 6. DASHBOARD DISPLAY
# ---------------------------------------------------------
m1, m2, m3, m4 = st.columns(4)
m1.metric("เส้นผ่านศูนย์กลางเทียบเท่า (dw)", f"{d_w_mm:.1f} mm ({d_w_cm:.2f} cm)")
m2.metric("ระยะอิทธิพลการระบาย (de)", f"{d_e_m:.2f} m ({d_e_cm:.0f} cm)")
m3.metric("การทรุดตัวสูงสุด (S final)", f"{S_final:.3f} m")
m4.metric("เวลาบรรลุ U = 90%", f"{days_90} วัน", 
          delta="ตามเป้าหมาย" if isinstance(days_90, (int, np.integer)) else "ช้าเกินไป", 
          delta_color="normal" if isinstance(days_90, (int, np.integer)) else "inverse")

st.markdown("<br>", unsafe_allow_html=True)

tab_steps, tab_charts, tab_data, tab_summary = st.tabs([
    "📑 วิธีทำแบบตาราง (Step-by-Step)", 
    "📊 กราฟวิเคราะห์ (Interactive Charts)", 
    "📋 ตารางข้อมูล (Data Table)", 
    "💡 สรุปผลการออกแบบ & Export Report"
])

# =========================================================
# TAB 1: ตารางแสดงวิธีทำทีละขั้นตอน
# =========================================================
with tab_steps:
    st.markdown(f"### 📑 ตารางคำนวณและตรวจสอบที่ระยะ $S = {S:.2f}$ m ณ เวลา $t = {target_day}$ วัน")
    st.caption("จำลองวิธีคิดเลขทีละคอลัมน์ ถอดแบบจากสไลด์เลคเชอร์และตำราปฐพีวิศวกรรม")
    
    # STEP 5: F(n)
    st.markdown("---")
    st.markdown("#### ⑤ ตารางคำนวณเพื่อหาค่า Drain Spacing Factor: $F(n)$")
    
    s_test_list = [round(max(0.5, S - 0.1), 2), round(S, 2), round(S + 0.1, 2)]
    fn_rows = []
    
    for s_val in s_test_list:
        de_val_m = 1.05 * s_val if "สามเหลี่ยม" in pattern else 1.13 * s_val
        de_val_cm = de_val_m * 100.0
        n_val = de_val_cm / d_w_cm
        n2 = n_val**2
        t1 = n2 / (n2 - 1)
        t2 = np.log(n_val)
        t3 = (3 * n2 - 1) / (4 * n2)
        fn_val = (t1 * t2) - t3
        
        fn_rows.append({
            "(1) S (m)": f"{s_val:.2f}",
            "(2) de (cm)": f"{de_val_cm:.1f}",
            "(3) n = de/dw": f"{n_val:.2f}",
            "(4) n²": f"{n2:.2f}",
            "(5) n²/(n²-1)": f"{t1:.3f}",
            "(6) ln(n)": f"{t2:.3f}",
            "(7) (3n²-1)/(4n²)": f"{t3:.3f}",
            "(8) F(n) = (5)×(6)-(7)": f"★ {fn_val:.3f}" if s_val == S else f"{fn_val:.3f}"
        })
    st.table(pd.DataFrame(fn_rows))
    st.info(f"💡 **ผลลัพธ์ข้อ ⑤:** ที่ระยะติดตั้ง $S = {S:.2f}$ m คำนวณค่าอัตราส่วน $n = {n:.2f}$ ได้ค่า **$F(n) = {Fn:.3f}$**")

    # STEP 6: Ur
    st.markdown("---")
    st.markdown("#### ⑥ ตารางคำนวณเพื่อหาค่า Degree of Consolidation ในแนวรัศมี ($U_r$)")
    
    de2_cm2 = d_e_cm**2
    test_days = sorted(list(set([30, 60, target_day, 180])))
    ur_rows = []
    
    for d_val in test_days:
        cr_t = Cr_cm2_day * d_val
        Tr_val = cr_t / de2_cm2
        exp_term = np.exp((-8 * Tr_val) / Fn)
        ur_val = (1 - exp_term) * 100.0
        
        ur_rows.append({
            "(1) S (m)": f"{S:.2f}",
            "(2) Cr (cm²/day)": f"{Cr_cm2_day:.1f}",
            "(3) t (day)": f"{d_val}",
            "(4) de² (cm²)": f"{de2_cm2:.0f}",
            "(5) Cr × t (cm²)": f"{cr_t:.0f}",
            "(6) Tr = (5)/(4)": f"{Tr_val:.4f}",
            "(7) F(n)": f"{Fn:.3f}",
            "(8) exp(-8Tr/Fn)": f"{exp_term:.4f}",
            "(9) Ur (%) = 1 - (8)": f"★ {ur_val:.2f}%" if d_val == target_day else f"{ur_val:.2f}%"
        })
    st.table(pd.DataFrame(ur_rows))
    
    Ur_target = df.loc[df["Day"] == target_day, "U_r"].values[0] if target_day <= 365 else (1 - np.exp((-8 * (Cr_cm2_day * target_day) / de2_cm2) / Fn)) * 100

    # STEP 7: Uv
    st.markdown("---")
    st.markdown(f"#### ⑦ คำนวณระดับการอัดตัวคายน้ำในแนวดิ่ง ($U_v$) ที่เวลา $t = {target_day}$ วัน (Theory of Terzaghi)")
    
    col_v1, col_v2 = st.columns([1, 1.5])
    with col_v1:
        st.latex(r"T_v = \frac{C_v \times t}{(H_d)^2}")
        st.latex(r"U_v = \frac{\sqrt{4 \times T_v}}{\pi} \quad (\text{เมื่อ } U_v \le 60\%)")
        st.latex(r"U_v = 1 - 10^{-0.085 - 0.933 T_v} \quad (\text{เมื่อ } U_v > 60\%)")
        
    with col_v2:
        Tv_target = (Cv_cm2_day * target_day) / (H_d_cm**2)
        Uv_target = np.sqrt(4 * Tv_target) / np.pi if Tv_target <= 0.286 else 1 - (10**(-0.085 - 0.933 * Tv_target))
        Uv_target_pct = min(Uv_target * 100.0, 100.0)
        
        st.write(f"**แสดงการแทนค่าตัวเลข ณ เวลาเป้าหมาย $t = {target_day}$ วัน:**")
        st.write(f"- ค่าสัมประสิทธิ์การอัดตัว $C_v = \\mathbf{{{Cv_cm2_day:.1f}}}$ cm²/day")
        st.write(f"- ความหนาชั้นดิน $H = {H_soil}$ m ➔ ระยะระบายน้ำ $H_d = \\mathbf{{{H_d_cm:.0f}}}$ cm ({drainage_type.split(' ')[0]})")
        st.write(f"- คำนวณ Time Factor: $T_v = \\frac{{{Cv_cm2_day:.1f} \\times {target_day}}}{{({H_d_cm:.0f})^2}} = \\mathbf{{{Tv_target:.4f}}}$")
        st.write(f"- **ผลลัพธ์ข้อ ⑦ ได้ค่า $U_v$:** $\\mathbf{{{Uv_target_pct:.2f}\\%}}$")

    # STEP 8 & 9: Uav
    st.markdown("---")
    st.markdown("#### ⑧-⑨ คำนวณระดับการอัดตัวคายน้ำเฉลี่ยรวม ($U_{av}$) - Theory of Carillo (1942)")
    st.latex(r"U_{av} = 1 - (1 - U_r)(1 - U_v)")
    
    Uav_target_pct = 100.0 * (1.0 - (1.0 - (Ur_target/100.0)) * (1.0 - (Uv_target_pct/100.0)))
    
    st.write(f"**แทนค่าที่เวลา $t = {target_day}$ วัน:**")
    st.latex(f"U_{{av}} = 1 - (1 - {Ur_target/100:.4f})(1 - {Uv_target_pct/100:.4f}) = \\mathbf{{{Uav_target_pct:.2f}\\%}}")
    
    summary_rows = []
    for d_val in sorted(list(set([30, 60, target_day, 180, 270, 365]))):
        if d_val <= 365:
            r = df[df["Day"] == d_val].iloc[0]
            ur_p, uv_p, uav_p = r['U_r'], r['U_v'], r['U_av']
        else:
            ur_p = (1 - np.exp((-8 * (Cr_cm2_day * d_val) / de2_cm2) / Fn)) * 100
            tv_tmp = (Cv_cm2_day * d_val) / (H_d_cm**2)
            uv_tmp = np.sqrt(4 * tv_tmp) / np.pi if tv_tmp <= 0.286 else 1 - (10**(-0.085 - 0.933 * tv_tmp))
            uv_p = min(uv_tmp * 100, 100)
            uav_p = 100 * (1 - (1 - ur_p/100) * (1 - uv_p/100))
            
        status = "✅ ผ่านเกณฑ์ (> 90%)" if uav_p >= 90 else "⏳ ยังไม่ถึงเกณฑ์"
        summary_rows.append({
            "เวลา t (วัน)": f"{d_val} วัน",
            "U_r แนวรัศมี (ข้อ ⑥)": f"{ur_p:.2f}%",
            "U_v แนวดิ่ง (ข้อ ⑦)": f"{uv_p:.2f}%",
            "⑧ U_av รวม = 1-(1-Ur)(1-Uv)": f"★ {uav_p:.2f}%" if d_val == target_day else f"{uav_p:.2f}%",
            "⑨ ตรวจระดับ U_av > 90%": status
        })
    st.table(pd.DataFrame(summary_rows))
    
    if Uav_target_pct >= 90.0:
        st.success(f"🎯 **บทสรุปข้อ ⑨:** ที่ระยะติดตั้ง PVD **S = {S:.2f} เมตร** ชั้นดินสามารถอัดตัวคายน้ำได้ **{Uav_target_pct:.2f}%** ณ วันที่ {target_day} ซึ่ง **ผ่านเกณฑ์ (> 90%)** เป็นไปตามข้อกำหนดการออกแบบครับ")
    else:
        st.warning(f"⚠️ **บทสรุปข้อ ⑨:** ณ วันเป้าหมายที่ {target_day} วัน ชั้นดินอัดตัวคายน้ำได้เพียง **{Uav_target_pct:.2f}%** (ยังไม่ผ่านเกณฑ์ > 90%) แนะนำให้ลดระยะห่าง $S$ ลง หรือเพิ่มเวลาการรอคอยครับ")

# =========================================================
# TAB 2-4: กราฟ, ข้อมูล และส่งออก Word Report
# =========================================================
with tab_charts:
    col_left, col_right = st.columns(2)
    with col_left:
        fig_settle = go.Figure()
        # เส้นกราฟสี Rose Red สว่างสดใส
        fig_settle.add_trace(go.Scatter(x=df["Day"], y=df["Settlement"], mode='lines', name='Settlement (m)', line=dict(color='#F43F5E', width=3.5)))
        fig_settle.update_layout(
            title=dict(text="📉 กราฟพัฒนาการการทรุดตัวตามเวลา", font=dict(size=16, color="#0F172A")),
            xaxis_title="เวลา (วัน)", 
            yaxis_title="การทรุดตัว (เมตร)", 
            yaxis=dict(autorange="reversed", gridcolor="#F1F5F9"), 
            xaxis=dict(gridcolor="#F1F5F9"),
            template="plotly_white", 
            height=380,
            paper_bgcolor="rgba(0,0,0,0)",
            plot_bgcolor="#FFFFFF",
            margin=dict(l=20, r=20, t=50, b=20)
        )
        st.plotly_chart(fig_settle, use_container_width=True)

    with col_right:
        fig_u = go.Figure()
        # เส้นกราฟโทนสว่างตัดกันชัดเจน
        fig_u.add_trace(go.Scatter(x=df["Day"], y=df["U_av"], name='รวม (U_av)', line=dict(color='#10B981', width=3.5)))
        fig_u.add_trace(go.Scatter(x=df["Day"], y=df["U_r"], name='แนวรัศมี PVD (U_r)', line=dict(color='#0EA5E9', width=2.5, dash='dash')))
        fig_u.add_trace(go.Scatter(x=df["Day"], y=df["U_v"], name='แนวดิ่ง ดิน (U_v)', line=dict(color='#8B5CF6', width=2, dash='dot')))
        fig_u.add_hline(y=90, line_dash="dash", line_color="#F59E0B", annotation_text="Target 90%")
        fig_u.update_layout(
            title=dict(text="📊 อัตราการอัดตัวคายน้ำ (Degree of Consolidation)", font=dict(size=16, color="#0F172A")),
            xaxis_title="เวลา (วัน)", 
            yaxis_title="Consolidation (%)", 
            yaxis=dict(range=[0, 105], gridcolor="#F1F5F9"), 
            xaxis=dict(gridcolor="#F1F5F9"),
            template="plotly_white", 
            height=380,
            paper_bgcolor="rgba(0,0,0,0)",
            plot_bgcolor="#FFFFFF",
            margin=dict(l=20, r=20, t=50, b=20)
        )
        st.plotly_chart(fig_u, use_container_width=True)

    with col_right:
        fig_u = go.Figure()
        fig_u.add_trace(go.Scatter(x=df["Day"], y=df["U_av"], name='รวม (U_av)', line=dict(color='#10B981', width=3)))
        fig_u.add_trace(go.Scatter(x=df["Day"], y=df["U_r"], name='แนวรัศมี PVD (U_r)', line=dict(color='#3B82F6', dash='dash')))
        fig_u.add_trace(go.Scatter(x=df["Day"], y=df["U_v"], name='แนวดิ่ง ดิน (U_v)', line=dict(color='#9CA3AF', dash='dot')))
        fig_u.add_hline(y=90, line_dash="dash", line_color="#F59E0B", annotation_text="Target 90%")
        fig_u.update_layout(title="📊 อัตราการอัดตัวคายน้ำ (Degree of Consolidation)", xaxis_title="เวลา (วัน)", yaxis_title="Consolidation (%)", yaxis=dict(range=[0, 105]), template="plotly_white", height=380)
        st.plotly_chart(fig_u, use_container_width=True)

with tab_data:
    st.subheader("ตารางแสดงค่าการคำนวณรายวัน")
    csv = df.to_csv(index=False).encode('utf-8')
    st.download_button("📥 ดาวน์โหลดข้อมูลเป็น CSV", data=csv, file_name="pvd_consolidation_result.csv", mime="text/csv")
    st.dataframe(df.style.format({"U_v": "{:.2f}%", "U_r": "{:.2f}%", "U_av": "{:.2f}%", "Settlement": "{:.4f} m"}), use_container_width=True, height=300)

with tab_summary:
    st.success(f"**สรุปผลการประเมิน:** ในระยะติดตั้ง PVD ที่ **{S:.2f} เมตร** รูปแบบ **{pattern}**")
    st.write(f"- การทรุดตัวทั้งหมดเมื่อสิ้นสุดกระบวนการ ($S_{{final}}$): **{S_final:.3f} เมตร**")
    if isinstance(days_90, (int, np.integer)):
        st.write(f"- ดินจะทรุดตัวถึง 90% ภายในเวลา **{days_90} วัน**")
    else:
        st.warning("- การติดตั้ง PVD ระยะนี้ยังไม่สามารถทำให้ดินทรุดตัวถึง 90% ได้ภายใน 1 ปี แนะนำให้ **ลดระยะห่าง S** ลง")

    st.markdown("---")
    st.markdown("#### 📄 ส่งออกรายงานสรุปผลพร้อมตารางคำนวณ (Word Document)")
    st.write("กดปุ่มด้านล่างเพื่อดาวน์โหลดไฟล์เอกสาร `.docx` สรุปผลพร้อมตารางคำนวณประกอบรายงาน")
    
    S_target = (Uav_target_pct / 100.0) * S_final
    
    doc_file = generate_word_report(
        pattern=pattern,
        S=S,
        H_soil=H_soil,
        H_pvd=H_pvd,
        Cv=Cv_cm2_day,
        Cr=Cr_cm2_day,
        target_day=target_day,
        Uav_pct=Uav_target_pct,
        Ur_pct=Ur_target,
        Uv_pct=Uv_target_pct,
        S_final=S_final,
        S_target=S_target,
        days_90=days_90,
        d_w_mm=d_w_mm,
        d_e_m=d_e_m,
        n=n,
        Fn=Fn,
        df=df
    )
    
    st.download_button(
        label="📥 ดาวน์โหลดรายงานสรุปผลพร้อมตาราง (Word File .docx)",
        data=doc_file,
        file_name=f"PVD_Engineering_Report_S{S}m_{target_day}days.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
