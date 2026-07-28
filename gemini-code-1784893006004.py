import streamlit as st
import numpy as np
import pandas as pd
import plotly.graph_objects as go
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------------------------------------------
# 1. PAGE CONFIGURATION & CUSTOM CSS
# ---------------------------------------------------------
st.set_page_config(
    page_title="PVD Settlement Analytics",
    page_icon="🏗️",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
    .main-header { font-size: 2.2rem; font-weight: 700; color: #1E3A8A; margin-bottom: 0.2rem; }
    .sub-header { font-size: 1rem; color: #6B7280; margin-bottom: 1.5rem; }
    div[data-testid="stMetricValue"] { font-size: 1.8rem; font-weight: bold; }
    .stTabs [data-baseweb="tab-list"] { gap: 8px; }
    .stTabs [data-baseweb="tab"] { padding-left: 16px; padding-right: 16px; border-radius: 4px; }
</style>
""", unsafe_allow_html=True)

# ---------------------------------------------------------
# 2. HELPER FUNCTION: สรุปผล + สร้างตารางใน WORD REPORT
# ---------------------------------------------------------
def generate_word_report(pattern, S, H_soil, H_sfinal, H_pvd, Cv, Cr, target_day, Uav_pct, Ur_pct, Uv_pct, S_final, S_target, days_90, d_w_mm, d_e_m, n, Fn, df):
    doc = Document()
    
    # 1. หัวข้อรายงาน
    title = doc.add_heading('รายงานสรุปผลการวิเคราะห์การเร่งการทรุดตัวด้วย PVD', level=1)
    if title.runs:
        title.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    doc.add_paragraph(f"วันที่ออกรายงาน: {pd.Timestamp.now().strftime('%d/%m/%Y')}")
    doc.add_paragraph("-" * 55)
    
    # 2. ข้อมูลพารามิเตอร์การออกแบบ
    doc.add_heading('1. ข้อมูลและพารามิเตอร์การออกแบบ (Input Parameters)', level=2)
    p1 = doc.add_paragraph()
    p1.add_run(f"ทำการติดตั้ง PVD ในชั้นดินอ่อนความหนาอัดตัว {H_soil:.1f} m (ความหนาคำนวณทรุดตัว S_final = {H_sfinal:.1f} m) โดยจัดวางในรูปแบบ {pattern} ที่ระยะห่าง S = {S:.2f} m รายละเอียดพารามิเตอร์สรุปดังตารางด้านล่าง:")

    # --- WORD TABLE 1: Input Parameters Table ---
    t1 = doc.add_table(rows=1, cols=3)
    t1.style = 'Table Grid'
    hdr_cells1 = t1.rows[0].cells
    hdr_cells1[0].text = 'พารามิเตอร์ (Parameter)'
    hdr_cells1[1].text = 'ค่าที่ใช้ (Value)'
    hdr_cells1[2].text = 'หน่วย (Unit)'
    
    for cell in hdr_cells1:
        cell.paragraphs[0].runs[0].font.bold = True
        
    params_data = [
        ("รูปแบบการจัดวาง (Pattern)", f"{pattern}", "-"),
        ("ระยะห่างการติดตั้ง (S)", f"{S:.2f}", "m"),
        ("ความหนาชั้นดินอ่อน (H_soil)", f"{H_soil:.1f}", "m"),
        ("ความหนาชั้นดินคำนวณทรุดตัว (H_sfinal)", f"{H_sfinal:.1f}", "m"),
        ("ความลึกแผ่น PVD (H_pvd)", f"{H_pvd:.1f}", "m"),
        ("ค่าสัมประสิทธิ์ Cv", f"{Cv:.1f}", "cm²/day"),
        ("ค่าสัมประสิทธิ์ Cr", f"{Cr:.1f}", "cm²/day"),
        ("ขนาดเส้นผ่านศูนย์กลาง dw / de", f"{d_w_mm:.1f} mm / {d_e_m:.2f} m", "-"),
        ("อัตราส่วน n (de/dw)", f"{n:.2f}", "-"),
        ("Drain Spacing Factor F(n)", f"{Fn:.3f}", "-"),
    ]
    
    for p_name, p_val, p_unit in params_data:
        row_cells = t1.add_row().cells
        row_cells[0].text = p_name
        row_cells[1].text = p_val
        row_cells[2].text = p_unit
        
    doc.add_paragraph() # เว้นบรรทัด

    # 3. ผลการคำนวณ ณ วันเป้าหมาย
    doc.add_heading('2. ผลการคำนวณ ณ วันเป้าหมาย (Target Day Results)', level=2)
    status_text = "บรรลุตามข้อกำหนดการออกแบบ (>= 90%)" if Uav_pct >= 90 else "ยังไม่บรรลุตามเป้าหมาย (< 90%)"
    
    p2 = doc.add_paragraph()
    p2.add_run(f"ณ ระยะเวลาเป้าหมายที่ ")
    p2.add_run(f"{target_day} วัน ").bold = True
    p2.add_run("ผลการคำนวณอัตราการอัดตัวคายน้ำและการทรุดตัว สรุปได้ดังตารางด้านล่าง:")

    # --- WORD TABLE 2: Target Day Summary Table ---
    t2 = doc.add_table(rows=1, cols=3)
    t2.style = 'Table Grid'
    hdr_cells2 = t2.rows[0].cells
    hdr_cells2[0].text = 'รายการคำนวณ'
    hdr_cells2[1].text = 'ผลการคำนวณ'
    hdr_cells2[2].text = 'หมายเหตุ / เกณฑ์'

    for cell in hdr_cells2:
        cell.paragraphs[0].runs[0].font.bold = True

    results_data = [
        ("Degree of Consolidation แนวดิ่ง (Uv)", f"{Uv_pct:.2f}%", "Terzaghi Theory"),
        ("Degree of Consolidation แนวรัศมี (Ur)", f"{Ur_pct:.2f}%", "Barron / Hansbo"),
        ("Degree of Consolidation เฉลี่ยรวม (Uav)", f"{Uav_pct:.2f}%", status_text),
        ("การทรุดตัวที่เกิดขึ้นแล้ว (St)", f"{S_target:.3f} m", f"จาก S_final = {S_final:.3f} m"),
        ("เวลาบรรลุ U = 90%", f"{days_90} วัน", "เป้าหมายมาตรฐาน")
    ]

    for r_name, r_val, r_note in results_data:
        row_cells = t2.add_row().cells
        row_cells[0].text = r_name
        row_cells[1].text = r_val
        row_cells[2].text = r_note

    doc.add_paragraph() # เว้นบรรทัด

    # 4. ตารางพัฒนาการตามช่วงเวลา (Progress Table)
    doc.add_heading('3. ตารางพัฒนาการอัดตัวคายน้ำและการทรุดตัวตามช่วงเวลา', level=2)
    
    # --- WORD TABLE 3: Consolidation Progress Over Time Table ---
    t3 = doc.add_table(rows=1, cols=6)
    t3.style = 'Table Grid'
    hdr_cells3 = t3.rows[0].cells
    hdr_cells3[0].text = 'เวลา t (วัน)'
    hdr_cells3[1].text = 'U_r (%)'
    hdr_cells3[2].text = 'U_v (%)'
    hdr_cells3[3].text = 'U_av (%)'
    hdr_cells3[4].text = 'การทรุดตัว S_t (m)'
    hdr_cells3[5].text = 'สถานะ (> 90%)'

    for cell in hdr_cells3:
        cell.paragraphs[0].runs[0].font.bold = True

    test_days = sorted(list(set([30, 60, target_day, 180, 270, 365])))
    for d_val in test_days:
        if d_val <= 365:
            r = df[df["Day"] == d_val].iloc[0]
            ur_p, uv_p, uav_p, st_p = r['U_r'], r['U_v'], r['U_av'], r['Settlement']
            st_status = "ผ่านเกณฑ์" if uav_p >= 90 else "ยังไม่ถึงเกณฑ์"
        else:
            ur_p = (1 - np.exp((-8 * (Cr * d_val) / (d_e_m * 100)**2) / Fn)) * 100
            tv_tmp = (Cv * d_val) / ((H_soil*100/2)**2)
            uv_tmp = np.sqrt(4 * tv_tmp) / np.pi if tv_tmp <= 0.286 else 1 - (10**(-0.085 - 0.933 * tv_tmp))
            uv_p = min(uv_tmp * 100, 100)
            uav_p = 100 * (1 - (1 - ur_p/100) * (1 - uv_p/100))
            st_p = (uav_p / 100.0) * S_final
            st_status = "ผ่านเกณฑ์" if uav_p >= 90 else "ยังไม่ถึงเกณฑ์"

        row_cells = t3.add_row().cells
        row_cells[0].text = f"{d_val} วัน"
        row_cells[1].text = f"{ur_p:.2f}%"
        row_cells[2].text = f"{uv_p:.2f}%"
        row_cells[3].text = f"{uav_p:.2f}%"
        row_cells[4].text = f"{st_p:.3f} m"
        row_cells[5].text = st_status

    doc.add_paragraph() # เว้นบรรทัด

    # 5. สรุปผลเชิงวิศวกรรม
    doc.add_heading('4. สรุปผลและข้อเสนอแนะเชิงวิศวกรรม', level=2)
    p3 = doc.add_paragraph()
    
    if Uav_pct >= 90.0:
        p3.add_run(
            f"สรุปการประเมิน: การติดตั้ง PVD ระยะห่าง {S:.2f} เมตร มีความเหมาะสมและเพียงพอทางการวิศวกรรม "
            f"โดยชั้นดินสามารถคายน้ำและเกิดการทรุดตัวได้ถึง {Uav_pct:.2f}% ภายในระยะเวลา {target_day} วัน "
            f"(ใช้เวลาเพียง {days_90} วันในการบรรลุ U=90%) ซึ่งเป็นไปตามมาตรฐานการควบคุมการทรุดตัวก่อนการเปิดใช้งานพื้นที่"
        )
    else:
        p3.add_run(
            f"ข้อสังเกตทางการวิศวกรรม: การติดตั้ง PVD ที่ระยะห่าง {S:.2f} เมตร ณ เวลา {target_day} วัน "
            f"ทำให้เกิดการอัดตัวคายน้ำได้เพียง {Uav_pct:.2f}% ซึ่งยังไม่ถึงเกณฑ์ 90% "
            f"จึงมีข้อเสนอแนะให้ปรับลดระยะห่างการติดตั้ง PVD (S) ให้แคบลง เช่น ปรับลดเหลือ {max(0.6, S-0.2):.2f} เมตร "
            f"หรือขยายระยะเวลาในการรอคอยการทรุดตัว (Preloading Period) เพิ่มเติม"
        )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ---------------------------------------------------------
# 3. HEADER
# ---------------------------------------------------------
st.markdown('<div class="main-header">🏗️ PVD Consolidation Analytics</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-header">ระบบจำลองการทรุดตัวและคำนวณเวลาการอัดตัวคายน้ำชั้นดินอ่อน (ปรับหน่วยตามตำราเรียน / สไลด์เลคเชอร์)</div>', unsafe_allow_html=True)

# ---------------------------------------------------------
# 4. SIDEBAR - INPUT PARAMETERS (ระบบจำค่าอัตโนมัติผ่าน URL)
# ---------------------------------------------------------
st.sidebar.title("⚙️ ตั้งค่าพารามิเตอร์การออกแบบ")

def get_param(key, default_type, default_val):
    val = st.query_params.get(key, default_val)
    try:
        if default_type == float: return float(val)
        if default_type == int: return int(val)
        if default_type == bool: return str(val).lower() == "true"
        return str(val)
    except:
        return default_val

def update_url():
    for key, val in st.session_state.items():
        if key.startswith("mem_"):
            param_name = key.replace("mem_", "")
            st.query_params[param_name] = str(val)

with st.sidebar.expander("📐 1. รูปแบบ PVD & ระยะติดตั้ง", expanded=True):
    pattern_options = ["สามเหลี่ยม (Triangular)", "สี่เหลี่ยม (Square)"]
    def_pat = get_param("pattern", str, pattern_options[0])
    pattern = st.selectbox("รูปแบบการจัดวาง", pattern_options, 
                           index=pattern_options.index(def_pat) if def_pat in pattern_options else 0, 
                           key="mem_pattern", on_change=update_url)
    
    S = st.number_input("ระยะห่างการติดตั้ง S (m)", value=get_param("S", float, 1.00), step=0.05, format="%.2f", key="mem_S", on_change=update_url)
    H_pvd = st.number_input("ความลึกแผ่น PVD (m)", value=get_param("H_pvd", float, 10.0), step=0.5, key="mem_H_pvd", on_change=update_url)
    a_mm = st.number_input("ความกว้าง PVD: a (mm)", value=get_param("a_mm", float, 100.0), step=5.0, key="mem_a_mm", on_change=update_url)
    b_mm = st.number_input("ความหนา PVD: b (mm)", value=get_param("b_mm", float, 5.0), step=0.5, key="mem_b_mm", on_change=update_url)

with st.sidebar.expander("🧪 2. คุณสมบัติชั้นดิน & ทางระบายน้ำ", expanded=True):
    H_soil = st.number_input("ความหนาชั้นดินอ่อน: H_soil (m)", value=get_param("H_soil", float, 30.0), step=1.0, key="mem_H_soil", on_change=update_url)
    
    drain_options = ["ระบายน้ำ 2 ทาง (บน-ล่าง: Hd = H/2)", "ระบายน้ำทางเดียว (Hd = H)"]
    def_drain = get_param("drainage", str, drain_options[0])
    drainage_type = st.radio("ลักษณะการระบายน้ำแนวดิ่ง (Drainage Path)", drain_options, 
                             index=drain_options.index(def_drain) if def_drain in drain_options else 0, 
                             key="mem_drainage", on_change=update_url)
    
    st.markdown("---")
    unit_options = ["cm²/day (ตรงตามสไลด์)", "m²/year"]
    def_unit = get_param("unit", str, unit_options[0])
    unit_choice = st.selectbox("เลือกหน่วยของ Cv และ Cr", unit_options, 
                               index=unit_options.index(def_unit) if def_unit in unit_options else 0, 
                               key="mem_unit", on_change=update_url)
    
    if "cm²/day" in unit_choice:
        Cv_input = st.number_input("ค่า Cv (cm²/day)", value=get_param("Cv", float, 20.0), step=1.0, key="mem_Cv", on_change=update_url)
        Cr_input = st.number_input("ค่า Cr (cm²/day)", value=get_param("Cr", float, 140.0), step=5.0, key="mem_Cr", on_change=update_url)
        Cv_m2_yr = (Cv_input / 10000.0) * 365.25
        Cr_m2_yr = (Cr_input / 10000.0) * 365.25
        Cv_cm2_day = Cv_input
        Cr_cm2_day = Cr_input
    else:
        Cv_m2_yr = st.number_input("ค่า Cv (m²/year)", value=get_param("Cv", float, 2.0), step=0.1, key="mem_Cv", on_change=update_url)
        Cr_m2_yr = st.number_input("ค่า Cr (m²/year)", value=get_param("Cr", float, 6.0), step=0.1, key="mem_Cr", on_change=update_url)
        Cv_cm2_day = (Cv_m2_yr / 365.25) * 10000.0
        Cr_cm2_day = (Cr_m2_yr / 365.25) * 10000.0

    st.markdown("---")
    use_lab_data = st.checkbox("คำนวณ e₀ จากผลทดสอบ Lab", value=get_param("uselab", bool, False), key="mem_uselab", on_change=update_url)
    if use_lab_data:
        Gs = st.number_input("Specific Gravity (Gs)", value=get_param("Gs", float, 2.65), step=0.01, key="mem_Gs", on_change=update_url)
        w_wet = st.number_input("น้ำหนักดินเปียก W_wet (g)", value=get_param("wwet", float, 120.0), step=1.0, key="mem_wwet", on_change=update_url)
        w_dry = st.number_input("น้ำหนักดินแห้ง W_dry (g)", value=get_param("wdry", float, 80.0), step=1.0, key="mem_wdry", on_change=update_url)
        V_cm3 = st.number_input("ปริมาตรวงแหวน V (cm³)", value=get_param("Vcm3", float, 60.0), step=1.0, key="mem_Vcm3", on_change=update_url)
        
        w_water_content = (w_wet - w_dry) / w_dry
        gamma_bulk = w_wet / V_cm3
        gamma_d = gamma_bulk / (1.0 + w_water_content)
        e0 = ((Gs * 1.0) / gamma_d) - 1.0
        st.info(f"💡 e₀ ที่คำนวณได้ = **{e0:.3f}**")
    else:
        e0 = st.number_input("Initial Void Ratio (e0)", value=get_param("e0", float, 2.00), step=0.1, key="mem_e0", on_change=update_url)

    Cc = st.number_input("Compression Index (Cc)", value=get_param("Cc", float, 0.80), step=0.05, key="mem_Cc", on_change=update_url)
    
    # ความหนาชั้นดินสำหรับ S_final ไว้บริเวณเดียวกับ e0 และ Cc ตามที่คุณต้องการ
    H_sfinal = st.number_input("ความหนาชั้นดินคำนวณทรุดตัว: H_sfinal (m)", value=get_param("H_sfinal", float, 30.0), step=1.0, key="mem_H_sfinal", on_change=update_url)

    sigma_0 = st.number_input("Effective Stress เดิม: σ0' (kPa)", value=get_param("sigma_0", float, 50.0), step=5.0, key="mem_sigma_0", on_change=update_url)
    delta_sigma = st.number_input("น้ำหนักถมเพิ่ม: Δσ (kPa)", value=get_param("delta_sigma", float, 80.0), step=5.0, key="mem_delta_sigma", on_change=update_url)

with st.sidebar.expander("🎯 3. กำหนดวันเป้าหมาย & Smear Effect", expanded=True):
    target_day = st.number_input("วันเป้าหมายหลักที่ต้องการตรวจสอบ (วัน)", value=get_param("target_day", int, 90), step=10, min_value=1, key="mem_target_day", on_change=update_url)
    include_smear = st.checkbox("คิดผลกระทบ Smear Effect", value=get_param("smear", bool, False), key="mem_smear", on_change=update_url)
    if include_smear:
        d_s_ratio = st.number_input("อัตราส่วน ds / dw", value=get_param("ds_ratio", float, 2.5), step=0.1, key="mem_ds_ratio", on_change=update_url)
        kh_ks_ratio = st.number_input("อัตราส่วน kh / ks", value=get_param("kh_ks", float, 3.0), step=0.5, key="mem_kh_ks", on_change=update_url)
    else:
        d_s_ratio, kh_ks_ratio = 1.0, 1.0

with st.sidebar.expander("🏖️ 4. Sand Mat (ชั้นทรายซับน้ำ)", expanded=True):
    include_sandmat = st.checkbox("คิดผลกระทบความต้านทาน Sand Mat", value=get_param("sandmat", bool, False), key="mem_sandmat", on_change=update_url)
    if include_sandmat:
        H_m = st.number_input("ความหนาแผ่นทราย H_m (m)", value=get_param("Hm", float, 0.80), step=0.1, key="mem_Hm", on_change=update_url)
        B_sand = st.number_input("ความกว้างครึ่งหนึ่งของแผ่นทราย B (m)", value=get_param("Bsand", float, 5.0), step=0.5, key="mem_Bsand", on_change=update_url)
        kc_value = st.number_input("การซึมน้ำดินเหนียว $k_c$ ($10^{-7}$ cm/s)", value=get_param("kc", float, 1e-7), format="%.2e", key="mem_kc", on_change=update_url)
        km_value = st.number_input("การซึมน้ำแผ่นทราย $k_m$ ($10^{-3}$ cm/s)", value=get_param("km", float, 1e-3), format="%.2e", key="mem_km", on_change=update_url)
    else:
        H_m, B_sand, kc_value, km_value = 0.80, 5.0, 1e-7, 1e-3

# =========================================================
# 5. CALCULATION ENGINE
# =========================================================
a_m, b_m = a_mm / 1000.0, b_mm / 1000.0
d_w_m = (a_m + b_m) / 2.0
d_w_cm = d_w_m * 100.0
d_w_mm = d_w_m * 1000.0

d_e_m = 1.05 * S if "สามเหลี่ยม" in pattern else 1.13 * S
d_e_cm = d_e_m * 100.0
n = d_e_m / d_w_m

Fn = ((n**2) / (n**2 - 1)) * np.log(n) - ((3 * n**2 - 1) / (4 * n**2))
F_n_eff = Fn + (kh_ks_ratio - 1) * np.log(d_s_ratio) if include_smear else Fn

if include_sandmat and H_m > 0 and km_value > 0:
    # L = (32 / pi^2) * (1 / n^2) * (H / H_m) * (kc / km) * (B / dw)^2  (ใช้ H_soil เป็นความลึก H ของชั้นดินเหนียว)
    L_factor = (32.0 / (np.pi ** 2)) * (1.0 / (n ** 2)) * (H_soil / H_m) * (kc_value / km_value) * ((B_sand / d_w_m) ** 2)
else:
    L_factor = 0.0

denom = F_n_eff + (0.8 * L_factor)
H_d_m = H_soil / 2.0 if "2 ทาง" in drainage_type else H_soil
H_d_cm = H_d_m * 100.0

# คำนวณ S_final โดยใช้ค่า H_sfinal ที่แยกออกมา
S_final = H_sfinal * (Cc / (1.0 + e0)) * np.log10((sigma_0 + delta_sigma) / sigma_0)

t_day = float(target_day)
Tv = (Cv_cm2_day * t_day) / (H_d_cm ** 2)
Uv = (2.0 * np.sqrt(Tv)) / np.sqrt(np.pi) if Tv <= 0.282 else 1.0 - (10 ** (-(Tv + 0.085) / 0.933))
Uv = min(max(Uv, 0.0), 1.0)

Tr = (Cr_cm2_day * t_day) / (d_e_cm ** 2)
Ur = 1.0 - np.exp((-8.0 * Tr) / denom)
Ur = min(max(Ur, 0.0), 1.0)

Uav = 1.0 - ((1.0 - Ur) * (1.0 - Uv))
St = Uav * S_final

days = np.arange(1, 366)
U_v_list, U_r_list, U_av_list, S_t_list = [], [], [], []

for d_day in days:
    Tv_i = (Cv_cm2_day * d_day) / (H_d_cm ** 2)
    U_v_i = (2.0 * np.sqrt(Tv_i)) / np.sqrt(np.pi) if Tv_i <= 0.282 else 1.0 - (10 ** (-(Tv_i + 0.085) / 0.933))
    U_v_i = min(max(U_v_i, 0.0), 1.0)
    
    Tr_i = (Cr_cm2_day * d_day) / (d_e_cm ** 2)
    U_r_i = 1.0 - np.exp((-8.0 * Tr_i) / denom)
    U_r_i = min(max(U_r_i, 0.0), 1.0)
    
    U_av_i = 1.0 - ((1.0 - U_r_i) * (1.0 - U_v_i))
    S_t_i = U_av_i * S_final
    
    U_v_list.append(U_v_i * 100.0)
    U_r_list.append(U_r_i * 100.0)
    U_av_list.append(U_av_i * 100.0)
    S_t_list.append(S_t_i)

df = pd.DataFrame({
    "Day": days, 
    "U_v": U_v_list, 
    "U_r": U_r_list, 
    "U_av": U_av_list, 
    "Settlement": S_t_list
})

u90_idx = df[df["U_av"] >= 90.0].first_valid_index()
days_90 = df.loc[u90_idx, "Day"] if u90_idx is not None else "> 365"

# ---------------------------------------------------------
# 6. DASHBOARD DISPLAY
# ---------------------------------------------------------
m1, m2, m3, m4 = st.columns(4)
m1.metric("เส้นผ่านศูนย์กลางเทียบเท่า (dw)", f"{d_w_mm:.1f} mm ({d_w_cm:.2f} cm)")
m2.metric("ระยะอิทธิพลการระบาย (de)", f"{d_e_m:.2f} m ({d_e_cm:.0f} cm)")
m3.metric("การทรุดตัวสูงสุด (S final)", f"{S_final:.3f} m")
m4.metric("เวลาบรรลุ U = 90%", f"{days_90} วัน", 
          delta="ตามเป้าหมาย" if isinstance(days_90, (int, np.integer)) else "ช้าเกินไป", 
          delta_color="normal" if isinstance(days_90, (int, np.integer)) else "inverse")

st.markdown("<br>", unsafe_allow_html=True)

tab_steps, tab_charts, tab_data, tab_summary = st.tabs([
    "📑 วิธีทำแบบตาราง (Step-by-Step)", 
    "📊 กราฟวิเคราะห์ (Interactive Charts)", 
    "📋 ตารางข้อมูล (Data Table)", 
    "💡 สรุปผลการออกแบบ & Export Report"
])

# =========================================================
# TAB 1: ตารางแสดงวิธีทำทีละขั้นตอน
# =========================================================
with tab_steps:
    st.markdown(f"### 📑 ตารางคำนวณและตรวจสอบที่ระยะ $S = {S:.2f}$ m ณ เวลา $t = {target_day}$ วัน")
    st.caption("จำลองวิธีคิดเลขทีละคอลัมน์ ถอดแบบจากสไลด์เลคเชอร์และตำราปฐพีวิศวกรรม")
    
    # STEP 5: F(n)
    st.markdown("---")
    st.markdown("#### ⑤ ตารางคำนวณเพื่อหาค่า Drain Spacing Factor: $F(n)$")
    
    s_test_list = [round(max(0.5, S - 0.1), 2), round(S, 2), round(S + 0.1, 2)]
    fn_rows = []
    
    for s_val in s_test_list:
        de_val_m = 1.05 * s_val if "สามเหลี่ยม" in pattern else 1.13 * s_val
        de_val_cm = de_val_m * 100.0
        n_val = de_val_cm / d_w_cm
        n2 = n_val**2
        t1 = n2 / (n2 - 1)
        t2 = np.log(n_val)
        t3 = (3 * n2 - 1) / (4 * n2)
        fn_val = (t1 * t2) - t3
        
        fn_rows.append({
            "(1) S (m)": f"{s_val:.2f}",
            "(2) de (cm)": f"{de_val_cm:.1f}",
            "(3) n = de/dw": f"{n_val:.2f}",
            "(4) n²": f"{n2:.2f}",
            "(5) n²/(n²-1)": f"{t1:.3f}",
            "(6) ln(n)": f"{t2:.3f}",
            "(7) (3n²-1)/(4n²)": f"{t3:.3f}",
            "(8) F(n) = (5)×(6)-(7)": f"★ {fn_val:.3f}" if s_val == S else f"{fn_val:.3f}"
        })
    st.table(pd.DataFrame(fn_rows))
    st.info(f"💡 **ผลลัพธ์ข้อ ⑤:** ที่ระยะติดตั้ง $S = {S:.2f}$ m คำนวณค่าอัตราส่วน $n = {n:.2f}$ ได้ค่า **$F(n) = {Fn:.3f}$**")

    # STEP 6: Ur & Sand Mat (ตามรูปที่เพิ่มเข้ามา)
    st.markdown("---")
    st.markdown("#### ⑥ การคำนวณอัตราการอัดตัวคายน้ำ กรณีพิจารณาผลกระทบจากทรายซับน้ำ (Sand Mat)")
    
    if include_sandmat:
        st.write("สมการคำนวณอัตราการอัดตัวคายน้ำในแนวรัศมี ($U_r$) เมื่อพิจารณาความต้านทานของแผ่นทรายซับน้ำ:")
        st.latex(r"U_r = 1 - \exp\left( \frac{-8T_r}{F(n) + 0.8L} \right)")
        
        st.write("โดยมีสมการคำนวณหาค่าดัชนีความต้านทานต่อการระบายน้ำ ($L$) ของทรายซับน้ำดังนี้:")
        st.latex(r"L = \frac{32}{\pi^2} \cdot \frac{1}{n^2} \cdot \frac{H}{H_m} \cdot \frac{k_c}{k_m} \cdot \left(\frac{B}{d_w}\right)^2")
        
        st.markdown("##### 📌 รายละเอียดตัวแปรและการแทนค่าคำนวณค่า $L$:")
        st.write(f"- **$n$ (Drainage spacing factor / อัตราส่วนพื้นที่ระบายน้ำ):** `{n:.2f}` (➔ $n^2 = {n**2:.2f}$)")
        st.write(f"- **$H$ (ความลึกของชั้นดินเหนียว):** `{H_soil}` m")
        st.write(f"- **$H_m$ (ความหนาของแผ่นทราย Sand Mat):** `{H_m}` m")
        st.write(f"- **$B$ (ครึ่งหนึ่งของความกว้างแผ่นทราย):** `{B_sand}` m")
        st.write(f"- **$d_w$ (ขนาดเส้นผ่านศูนย์กลางเทียบเท่าของท่อระบายน้ำ):** `{d_w_m:.3f}` m")
        st.write(f"- **$k_c$ (ค่าการซึมน้ำของดินเหนียว):** `{kc_value:.1e}` cm/s")
        st.write(f"- **$k_m$ (ค่าการซึมน้ำของแผ่นทราย):** `{km_value:.1e}` cm/s")
        
        col_sm1, col_sm2 = st.columns(2)
        with col_sm1:
            st.success(f"➔ **ค่าดัชนีความต้านทาน $L$ ที่คำนวณได้:** `L = {L_factor:.4f}`")
        with col_sm2:
            st.info(f"➔ **ตัวหารรวมในสมการ $F(n) + 0.8L$:** `{Fn:.3f} + 0.8({L_factor:.4f}) = {denom:.4f}`")
    else:
        st.latex(r"U_r = 1 - \exp\left( \frac{-8T_r}{F(n)} \right)")
        st.info("ℹ️ ปัจจุบันไม่ได้เลือกใช้งานการคิดผลกระทบจาก Sand Mat (ค่าสัมประสิทธิ์ $L = 0$ และตัวหารใช้เพียง $F(n)$ ปกติ)")

    # ตารางคำนวณ Ur รายช่วงเวลา
    de2_cm2 = d_e_cm**2
    test_days = sorted(list(set([30, 60, target_day, 180])))
    ur_rows = []
    
    for d_val in test_days:
        cr_t = Cr_cm2_day * d_val
        Tr_val = cr_t / de2_cm2
        exp_term = np.exp((-8 * Tr_val) / denom)
        ur_val = (1 - exp_term) * 100.0
        
        ur_rows.append({
            "(1) S (m)": f"{S:.2f}",
            "(2) Cr (cm²/day)": f"{Cr_cm2_day:.1f}",
            "(3) t (day)": f"{d_val}",
            "(4) de² (cm²)": f"{de2_cm2:.0f}",
            "(5) Cr × t (cm²)": f"{cr_t:.0f}",
            "(6) Tr = (5)/(4)": f"{Tr_val:.4f}",
            "(7) ตัวหาร (Fn + 0.8L)": f"{denom:.3f}",
            "(8) exp(-8Tr/ตัวหาร)": f"{exp_term:.4f}",
            "(9) Ur (%) = 1 - (8)": f"★ {ur_val:.2f}%" if d_val == target_day else f"{ur_val:.2f}%"
        })
    st.table(pd.DataFrame(ur_rows))
    
    Ur_target = df.loc[df["Day"] == target_day, "U_r"].values[0] if target_day <= 365 else (1 - np.exp((-8 * (Cr_cm2_day * target_day) / de2_cm2) / denom)) * 100

    # STEP 7: Uv
    st.markdown("---")
    st.markdown(f"#### ⑦ คำนวณระดับการอัดตัวคายน้ำในแนวดิ่ง ($U_v$) ที่เวลา $t = {target_day}$ วัน (Theory of Terzaghi)")
    
    col_v1, col_v2 = st.columns([1, 1.5])
    with col_v1:
        st.latex(r"T_v = \frac{C_v \times t}{(H_d)^2}")
        st.latex(r"U_v = \frac{\sqrt{4 \times T_v}}{\pi} \quad (\text{เมื่อ } U_v \le 60\%)")
        st.latex(r"U_v = 1 - 10^{-0.085 - 0.933 T_v} \quad (\text{เมื่อ } U_v > 60\%)")
        
    with col_v2:
        Tv_target = (Cv_cm2_day * target_day) / (H_d_cm**2)
        Uv_target = np.sqrt(4 * Tv_target) / np.pi if Tv_target <= 0.286 else 1 - (10**(-0.085 - 0.933 * Tv_target))
        Uv_target_pct = min(Uv_target * 100.0, 100.0)
        
        st.write(f"**แสดงการแทนค่าตัวเลข ณ เวลาเป้าหมาย $t = {target_day}$ วัน:**")
        st.write(f"- ค่าสัมประสิทธิ์การอัดตัว $C_v = \\mathbf{{{Cv_cm2_day:.1f}}}$ cm²/day")
        st.write(f"- ความหนาชั้นดิน $H = {H_soil}$ m ➔ ระยะระบายน้ำ $H_d = \\mathbf{{{H_d_cm:.0f}}}$ cm ({drainage_type.split(' ')[0]})")
        st.write(f"- คำนวณ Time Factor: $T_v = \\frac{{{Cv_cm2_day:.1f} \\times {target_day}}}{{({H_d_cm:.0f})^2}} = \\mathbf{{{Tv_target:.4f}}}$")
        st.write(f"- **ผลลัพธ์ข้อ ⑦ ได้ค่า $U_v$:** $\\mathbf{{{Uv_target_pct:.2f}\\%}}$")

    # STEP 8 & 9: Uav
    st.markdown("---")
    st.markdown("#### ⑧-⑨ คำนวณระดับการอัดตัวคายน้ำเฉลี่ยรวม ($U_{av}$) - Theory of Carillo (1942)")
    st.latex(r"U_{av} = 1 - (1 - U_r)(1 - U_v)")
    
    Uav_target_pct = 100.0 * (1.0 - (1.0 - (Ur_target/100.0)) * (1.0 - (Uv_target_pct/100.0)))
    
    st.write(f"**แทนค่าที่เวลา $t = {target_day}$ วัน:**")
    st.latex(f"U_{{av}} = 1 - (1 - {Ur_target/100:.4f})(1 - {Uv_target_pct/100:.4f}) = \\mathbf{{{Uav_target_pct:.2f}\\%}}")
    
    summary_rows = []
    for d_val in sorted(list(set([30, 60, target_day, 180, 270, 365]))):
        if d_val <= 365:
            r = df[df["Day"] == d_val].iloc[0]
            ur_p, uv_p, uav_p = r['U_r'], r['U_v'], r['U_av']
        else:
            ur_p = (1 - np.exp((-8 * (Cr_cm2_day * d_val) / de2_cm2) / denom)) * 100
            tv_tmp = (Cv_cm2_day * d_val) / (H_d_cm**2)
            uv_tmp = np.sqrt(4 * tv_tmp) / np.pi if tv_tmp <= 0.286 else 1 - (10**(-0.085 - 0.933 * tv_tmp))
            uv_p = min(uv_tmp * 100, 100)
            uav_p = 100 * (1 - (1 - ur_p/100) * (1 - uv_p/100))
            
        status = "✅ ผ่านเกณฑ์ (> 90%)" if uav_p >= 90 else "⏳ ยังไม่ถึงเกณฑ์"
        summary_rows.append({
            "เวลา t (วัน)": f"{d_val} วัน",
            "U_r แนวรัศมี (ข้อ ⑥)": f"{ur_p:.2f}%",
            "U_v แนวดิ่ง (ข้อ ⑦)": f"{uv_p:.2f}%",
            "⑧ U_av รวม = 1-(1-Ur)(1-Uv)": f"★ {uav_p:.2f}%" if d_val == target_day else f"{uav_p:.2f}%",
            "⑨ ตรวจระดับ U_av > 90%": status
        })
    st.table(pd.DataFrame(summary_rows))
    
    if Uav_target_pct >= 90.0:
        st.success(f"🎯 **บทสรุปข้อ ⑨:** ที่ระยะติดตั้ง PVD **S = {S:.2f} เมตร** ชั้นดินสามารถอัดตัวคายน้ำได้ **{Uav_target_pct:.2f}%** ณ วันที่ {target_day} ซึ่ง **ผ่านเกณฑ์ (> 90%)** เป็นไปตามข้อกำหนดการออกแบบครับ")
    else:
        st.warning(f"⚠️ **บทสรุปข้อ ⑨:** ณ วันเป้าหมายที่ {target_day} วัน ชั้นดินอัดตัวคายน้ำได้เพียง **{Uav_target_pct:.2f}%** (ยังไม่ผ่านเกณฑ์ > 90%) แนะนำให้ลดระยะห่าง $S$ ลง หรือเพิ่มเวลาการรอคอยครับ")

    # ★ การคำนวณ S_final (ใช้อ้างอิงค่า H_sfinal ที่แยกออกมาใหม่)
    st.markdown("---")
    st.markdown("#### ★ การคำนวณการทรุดตัวสูงสุด ($S_{final}$) ที่เกิดขึ้นเมื่อสิ้นสุดกระบวนการ")
    col_sf1, col_sf2 = st.columns([1, 1.5])
    with col_sf1:
        st.latex(r"S_{final} = H_{sfinal} \cdot \frac{C_c}{1 + e_0} \cdot \log_{10}\left(\frac{\sigma_0' + \Delta\sigma}{\sigma_0'}\right)")
    with col_sf2:
        st.write("**แสดงการแทนค่าตัวเลขจากพารามิเตอร์ที่กำหนด:**")
        st.latex(f"S_{{final}} = {H_sfinal:.2f} \\cdot \\frac{{{Cc:.3f}}}{{1 + {e0:.3f}}} \\cdot \\log_{{10}}\\left(\\frac{{{sigma_0:.2f} + {delta_sigma:.2f}}}{{{sigma_0:.2f}}}\\right)")
        st.write(f"- 💡 **ความหนาชั้นดินคำนวณทรุดตัว ($H_{{sfinal}}$):** $\\mathbf{{{H_sfinal:.2f}}}$ **เมตร**")
        st.write(f"- 💡 **การทรุดตัวสูงสุด ($S_{{final}}$):** $\\mathbf{{{S_final:.4f}}}$ **เมตร** (หรือ `{S_final*100:.2f}` cm)")
        st.write(f"- 💡 **การทรุดตัว ณ วันที่ {target_day} ($S_t = U_{{av}} \\times S_{{final}}$):** $\\mathbf{{{St:.4f}}}$ **เมตร** (หรือ `{St*100:.2f}` cm)")

# =========================================================
# TAB 2-4: กราฟ, ข้อมูล และส่งออก Word Report
# =========================================================
with tab_charts:
    col_left, col_right = st.columns(2)
    with col_left:
        fig_settle = go.Figure()
        fig_settle.add_trace(go.Scatter(x=df["Day"], y=df["Settlement"], mode='lines', name='Settlement (m)', line=dict(color='#EF4444', width=3)))
        fig_settle.update_layout(title="📉 กราฟพัฒนาการการทรุดตัวตามเวลา", xaxis_title="เวลา (วัน)", yaxis_title="การทรุดตัว (เมตร)", yaxis=dict(autorange="reversed"), template="plotly_white", height=380)
        st.plotly_chart(fig_settle, use_container_width=True)

    with col_right:
        fig_u = go.Figure()
        fig_u.add_trace(go.Scatter(x=df["Day"], y=df["U_av"], name='รวม (U_av)', line=dict(color='#10B981', width=3)))
        fig_u.add_trace(go.Scatter(x=df["Day"], y=df["U_r"], name='แนวรัศมี PVD (U_r)', line=dict(color='#3B82F6', dash='dash')))
        fig_u.add_trace(go.Scatter(x=df["Day"], y=df["U_v"], name='แนวดิ่ง ดิน (U_v)', line=dict(color='#9CA3AF', dash='dot')))
        fig_u.add_hline(y=90, line_dash="dash", line_color="#F59E0B", annotation_text="Target 90%")
        fig_u.update_layout(title="📊 อัตราการอัดตัวคายน้ำ (Degree of Consolidation)", xaxis_title="เวลา (วัน)", yaxis_title="Consolidation (%)", yaxis=dict(range=[0, 105]), template="plotly_white", height=380)
        st.plotly_chart(fig_u, use_container_width=True)

with tab_data:
    st.subheader("ตารางแสดงค่าการคำนวณรายวัน")
    csv = df.to_csv(index=False).encode('utf-8')
    st.download_button("📥 ดาวน์โหลดข้อมูลเป็น CSV", data=csv, file_name="pvd_consolidation_result.csv", mime="text/csv")
    st.dataframe(df.style.format({"U_v": "{:.2f}%", "U_r": "{:.2f}%", "U_av": "{:.2f}%", "Settlement": "{:.4f} m"}), use_container_width=True, height=300)

with tab_summary:
    st.success(f"**สรุปผลการประเมิน:** ในระยะติดตั้ง PVD ที่ **{S:.2f} เมตร** รูปแบบ **{pattern}**")
    st.write(f"- การทรุดตัวทั้งหมดเมื่อสิ้นสุดกระบวนการ ($S_{{final}}$ จาก $H_{{sfinal}} = {H_sfinal:.1f}$ m): **{S_final:.3f} เมตร**")
    if isinstance(days_90, (int, np.integer)):
        st.write(f"- ดินจะทรุดตัวถึง 90% ภายในเวลา **{days_90} วัน**")
    else:
        st.warning("- การติดตั้ง PVD ระยะนี้ยังไม่สามารถทำให้ดินทรุดตัวถึง 90% ได้ภายใน 1 ปี แนะนำให้ **ลดระยะห่าง S** ลง")

    st.markdown("---")
    st.markdown("#### 📄 ส่งออกรายงานสรุปผลพร้อมตารางคำนวณ (Word Document)")
    st.write("กดปุ่มด้านล่างเพื่อดาวน์โหลดไฟล์เอกสาร `.docx` สรุปผลพร้อมตารางคำนวณประกอบรายงาน")
    
    Uav_target_pct = Uav * 100.0
    Ur_target = Ur * 100.0
    Uv_target_pct = Uv * 100.0
    S_target = St

    doc_file = generate_word_report(
        pattern=pattern,
        S=S,
        H_soil=H_soil,
        H_sfinal=H_sfinal,
        H_pvd=H_pvd,
        Cv=Cv_cm2_day,
        Cr=Cr_cm2_day,
        target_day=target_day,
        Uav_pct=Uav_target_pct,
        Ur_pct=Ur_target,
        Uv_pct=Uv_target_pct,
        S_final=S_final,
        S_target=S_target,
        days_90=days_90,
        d_w_mm=d_w_mm,
        d_e_m=d_e_m,
        n=n,
        Fn=Fn,
        df=df
    )
        
    st.download_button(
        label="📥 ดาวน์โหลดรายงานสรุปผลพร้อมตาราง (Word File .docx)",
        data=doc_file,
        file_name=f"PVD_Engineering_Report_S{S}m_{target_day}days.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
